import google.generativeai as genai
import docx
from docx.shared import Pt, Inches # Added Inches for potential use
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
import re
import copy # Needed for deep copying elements if attempting higher fidelity merge later

# --- Configuration ---

def configure_gemini(api_key):
    """Configures the Google Generative AI SDK."""
    try:
        genai.configure(api_key=api_key)
        return True
    except Exception as e:
        print(f"Error configuring Gemini: {e}")
        return False

# --- Document Parsing (Code remains the same as qna_backend_v4) ---

def parse_qna_pairs(docx_file):
    """
    Parses a DOCX file to extract Q&A pairs based on specific keywords.
    Primarily looks for paragraphs starting with 'السؤال'/'سؤال'/'Question'
    or 'الجواب'/'جواب'/'Answer'. Handles Arabic text.

    Args:
        docx_file (file-like object): The uploaded Word document.

    Returns:
        list: A list of dictionaries, each containing 'question', 'answer',
              and 'q_para_index' (index of the first paragraph of the question).
              Returns an empty list if parsing fails or no pairs are found.
        list: A list of all paragraphs from the original document.
    """
    qna_pairs = []
    current_q = None
    current_a = []
    q_start_index = -1
    is_parsing_answer = False

    try:
        # Ensure the file stream is read correctly
        # If docx_file is UploadedFile, its stream might need seeking
        if hasattr(docx_file, 'seek'):
            docx_file.seek(0)
        document = docx.Document(docx_file)
        all_paragraphs = document.paragraphs

        q_pattern = re.compile(r"^\s*(?:السؤال|سؤال|Question)\b\s*[:.)]?\s*", re.IGNORECASE)
        a_pattern = re.compile(r"^\s*(?:الجواب|جواب|Answer)\b\s*[:.)]?\s*", re.IGNORECASE)

        for i, para in enumerate(all_paragraphs):
            text = para.text.strip()
            if not text:
                continue

            is_question_start = q_pattern.match(text)
            is_answer_start = a_pattern.match(text)

            if is_question_start:
                if current_q and current_a:
                    qna_pairs.append({
                        "question": current_q.strip(),
                        "answer": "\n".join(current_a).strip(),
                        "q_para_index": q_start_index
                    })
                elif current_q and not current_a:
                     print(f"Warning: Question starting at index {q_start_index} ('{current_q[:50]}...') seems to have no corresponding answer before the next question.")

                current_q = q_pattern.sub('', text).strip()
                q_start_index = i
                current_a = []
                is_parsing_answer = False

            elif is_answer_start and current_q:
                answer_text = a_pattern.sub('', text).strip()
                if answer_text:
                    current_a.append(answer_text)
                is_parsing_answer = True

            elif current_q:
                 if is_parsing_answer:
                     if text: current_a.append(text)
                 else:
                     if text: current_q += "\n" + text

        if current_q and current_a:
            qna_pairs.append({
                "question": current_q.strip(),
                "answer": "\n".join(current_a).strip(),
                "q_para_index": q_start_index
            })
        elif current_q and not current_a:
             print(f"Warning: The last question found ('{current_q[:50]}...') appears to have no corresponding answer.")

        return qna_pairs, all_paragraphs

    except Exception as e:
        print(f"Error parsing DOCX: {e}")
        # Consider logging the exception traceback for debugging
        import traceback
        traceback.print_exc()
        return [], []

# --- Headline Generation (Code remains the same as qna_backend_v4) ---

def generate_headline(question, answer, model_name="gemini-1.5-flash"):
    """
    Generates a headline for a given Q&A pair using the Gemini API.
    (Kept the Arabic prompt as document content is Arabic)
    """
    prompt = f"""
    السؤال التالي وإجابته مقتبسان من وثيقة. قم بإنشاء عنوان قصير ومناسب باللغة العربية يلخص الموضوع الرئيسي لهذا السؤال والإجابة. اجعل العنوان موجزًا وواضحًا. لا تقم بتضمين الكلمات "سؤال" أو "إجابة" في العنوان.

    السؤال:
    {question}

    الإجابة:
    {answer}

    العنوان المقترح:
    """
    try:
        model = genai.GenerativeModel(model_name)
        # Add safety settings if needed, though defaults are usually fine
        # response = model.generate_content(prompt, safety_settings={'HARASSMENT': 'BLOCK_NONE'})
        response = model.generate_content(prompt)

        # Check for valid response and text attribute
        if hasattr(response, 'text') and response.text:
            headline = response.text.strip().replace('*', '').replace('#', '')
            return headline
        else:
            # Log potential blocking or empty response more clearly
            block_reason = "Unknown"
            if hasattr(response, 'prompt_feedback') and response.prompt_feedback:
                 block_reason = str(response.prompt_feedback)
            print(f"Warning: Received empty or potentially blocked response for Q: {question[:50]}... Reason: {block_reason}")
            return f"خطأ في إنشاء عنوان للسؤال: {question[:30]}... (استجابة فارغة أو محظورة)"

    except Exception as e:
        print(f"Error generating headline with Gemini: {e}")
        import traceback
        traceback.print_exc()
        return f"خطأ في إنشاء عنوان للسؤال: {question[:30]}... (خطأ API)"


# --- Document Creation (Code remains the same as qna_backend_v4) ---

def create_modified_document(original_paragraphs, qna_pairs_with_headlines):
    """
    Creates a new DOCX document object with headlines inserted before Q&A pairs.

    Args:
        original_paragraphs (list): List of all paragraphs from the original doc.
        qna_pairs_with_headlines (list): List of Q&A dicts, now including a 'headline' key.

    Returns:
        docx.Document: The new document object with headlines, or None on error.
    """
    try:
        new_document = docx.Document()
        # Set default paragraph direction to RTL for Arabic
        style = new_document.styles['Normal']
        font = style.font
        # font.name = 'Arial' # Uncomment and set appropriate font if needed
        paragraph_format = style.paragraph_format
        paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT # Default to right align

        headlines_map = {item['q_para_index']: item['headline'] for item in qna_pairs_with_headlines}
        processed_para_indices = set()
        qna_ranges = {} # Store start and end index for each Q&A pair

        # Pre-calculate ranges to handle copying correctly
        sorted_qna = sorted(qna_pairs_with_headlines, key=lambda x: x['q_para_index'])
        for i, item in enumerate(sorted_qna):
            q_index = item['q_para_index']
            next_q_start_index = len(original_paragraphs)
            if i + 1 < len(sorted_qna):
                next_q_start_index = sorted_qna[i+1]['q_para_index']

            qna_ranges[q_index] = next_q_start_index
            for j in range(q_index, next_q_start_index):
                processed_para_indices.add(j)

        para_index = 0
        while para_index < len(original_paragraphs):
            if para_index in headlines_map:
                # Insert headline
                headline_text = headlines_map[para_index]
                headline_para = new_document.add_paragraph()
                headline_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
                runner = headline_para.add_run(headline_text)
                runner.bold = True
                runner.font.size = Pt(14) # Optional: make headline slightly larger

                end_index = qna_ranges[para_index]

                # Copy the original Q&A paragraphs for this range
                for i in range(para_index, end_index):
                    original_para = original_paragraphs[i]
                    # Skip empty paragraphs when copying if desired (optional)
                    # if not original_para.text.strip(): continue
                    new_para = new_document.add_paragraph(original_para.text)
                    # Attempt to copy basic formatting (alignment)
                    new_para.alignment = original_para.alignment # Copy alignment
                    # Note: Copying detailed formatting requires more complex run iteration.

                para_index = end_index

            elif para_index not in processed_para_indices:
                 # Copy paragraphs that are not part of any detected Q&A
                 original_para = original_paragraphs[para_index]
                 # Skip empty paragraphs when copying if desired (optional)
                 # if not original_para.text.strip(): continue
                 new_para = new_document.add_paragraph(original_para.text)
                 new_para.alignment = original_para.alignment
                 para_index += 1
            else:
                para_index += 1

        return new_document
    except Exception as e:
        print(f"Error creating modified document structure: {e}")
        import traceback
        traceback.print_exc()
        return None # Return None to indicate failure


# --- NEW: Document Merging ---

def merge_documents(doc_list):
    """
    Merges multiple docx.Document objects into a single document.

    Args:
        doc_list (list): A list of python-docx Document objects.

    Returns:
        docx.Document: A new Document object containing the merged content.

    Note: This function performs a basic content merge. It copies paragraph text
          and basic run formatting (bold, italic, size). It does NOT preserve
          headers, footers, sections, complex table formatting, or document properties.
          A page break is inserted between the content of each original document.
    """
    merged_document = docx.Document()
    # Set default style alignment to Right for Arabic content
    style = merged_document.styles['Normal']
    paragraph_format = style.paragraph_format
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    first_doc = True
    for doc_to_merge in doc_list:
        if not doc_to_merge: # Skip if a document object is None (e.g., due to processing error)
            print("Warning: Skipping a None document object during merge.")
            continue

        # Add a page break before appending subsequent documents
        if not first_doc:
            merged_document.add_page_break()
        first_doc = False

        # Iterate through paragraphs in the document to merge
        for para in doc_to_merge.paragraphs:
            new_para = merged_document.add_paragraph()
            # Copy paragraph alignment
            new_para.alignment = para.alignment
            # Copy paragraph style if needed (can be complex)
            # new_para.style = para.style

            # Copy text and basic formatting by iterating through runs
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                # Copy basic formatting
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.size = run.font.size
                new_run.font.name = run.font.name
                # Copy color if needed (requires RGBColor import)
                # from docx.shared import RGBColor
                # if run.font.color.rgb:
                #    new_run.font.color.rgb = run.font.color.rgb

        # Basic table copying (copies text content, not complex formatting)
        # More robust table copying is significantly more complex
        # for table in doc_to_merge.tables:
        #     new_table = merged_document.add_table(rows=len(table.rows), cols=len(table.columns))
        #     # Consider adding style: new_table.style = 'Table Grid'
        #     for i, row in enumerate(table.rows):
        #         for j, cell in enumerate(row.cells):
        #             new_table.cell(i, j).text = cell.text


    return merged_document

# --- Saving Document (Code remains the same as qna_backend_v4) ---

def save_doc_to_bytes(document):
    """Saves the DOCX document object to a byte stream."""
    doc_io = io.BytesIO()
    document.save(doc_io)
    doc_io.seek(0) # Rewind the stream to the beginning
    return doc_io.getvalue()

